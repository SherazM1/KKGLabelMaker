"""Parser service for DOCX Shipment Request Form BOL uploads."""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Any

from docx import Document

from app.models.bol_standard_record import (
    BolAddressBlock,
    BolStandardItemLine,
    BolStandardRecord,
)


FIELD_ALIASES: dict[str, tuple[str, ...]] = {
    "origin_facility": ("Origin Facility",),
    "origin_street_address": ("Origin Street Address",),
    "origin_city_state_zip": ("Origin City/State/Zip", "Origin City State Zip"),
    "delivery_facility": ("Delivery Facility",),
    "delivery_street_address": ("Delivery Street Address",),
    "delivery_city_state_zip": ("Delivery City/State/Zip", "Delivery City State Zip"),
    "pallet_qty": ("Pallet Qty", "Pallet Quantity"),
    "pallet_dims": ("Pallet DIMS", "Pallet Dims", "Pallet Dimensions"),
    "pallet_weight": ("Pallet Weight",),
    "delivery_number": ("Delivery # (If Applicable)", "Delivery #", "Delivery Number"),
    "project": ("Project",),
    "comments": ("Comments",),
    "completed_by_logistics_team": ("Completed by Logistics Team",),
    "kkg_po": ("KKG PO#", "KKG PO #", "KKG PO"),
    "kkg_load": ("KKG Load #", "KKG Load#", "KKG Load"),
    "carrier_pro": ("Carrier Pro", "Carrier PRO", "Carrier Pro #"),
    "carrier_scac": ("Carrier SCAC", "SCAC"),
    "item_number": ("Item #", "ITEM #", "Item Number", "ITEM Number"),
    "upc": ("UPC", "UPC #"),
}


@dataclass(slots=True)
class BolDocUploadParseResult:
    """Parsed DOCX upload payload for Standard-family BOL generation."""

    fields: dict[str, str]
    records: list[BolStandardRecord]
    readable_text: str


def _normalize_label(value: str) -> str:
    cleaned = str(value or "").strip().lower()
    cleaned = re.sub(r"\([^)]*\)", " ", cleaned)
    cleaned = cleaned.replace("#", " number ")
    cleaned = re.sub(r"[^a-z0-9]+", " ", cleaned)
    return re.sub(r"\s+", " ", cleaned).strip()


ALIAS_LOOKUP: dict[str, str] = {
    _normalize_label(alias): field_name
    for field_name, aliases in FIELD_ALIASES.items()
    for alias in aliases
}


def _clean_text(value: str) -> str:
    cleaned = str(value or "").replace("\xa0", " ")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    return cleaned.strip()


def _extract_document_lines(doc: Document) -> list[str]:
    lines: list[str] = []

    for paragraph in doc.paragraphs:
        text = _clean_text(paragraph.text)
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cell_values = [_clean_text(cell.text) for cell in row.cells]
            for cell_value in cell_values:
                if cell_value:
                    lines.append(cell_value)
            non_empty_cells = [cell_value for cell_value in cell_values if cell_value]
            if len(non_empty_cells) >= 2:
                lines.append(f"{non_empty_cells[0]}: {non_empty_cells[1]}")

    return lines


def _split_label_value(line: str) -> tuple[str, str] | None:
    for separator in (":", "\t"):
        if separator in line:
            label, value = line.split(separator, 1)
            return _clean_text(label), _clean_text(value)

    normalized_line = _normalize_label(line)
    if normalized_line in ALIAS_LOOKUP:
        return None

    for alias in sorted(ALIAS_LOOKUP, key=len, reverse=True):
        if not normalized_line.startswith(alias):
            continue
        raw_alias = next(
            (
                candidate_alias
                for aliases in FIELD_ALIASES.values()
                for candidate_alias in aliases
                if _normalize_label(candidate_alias) == alias
            ),
            "",
        )
        value = _clean_text(line[len(raw_alias) :])
        if value:
            return raw_alias, value

    return None


def _extract_fields_from_lines(lines: list[str]) -> dict[str, str]:
    fields: dict[str, str] = {}

    for index, line in enumerate(lines):
        split_value = _split_label_value(line)
        if split_value is not None:
            label, value = split_value
            field_name = ALIAS_LOOKUP.get(_normalize_label(label))
            if field_name and value and field_name not in fields:
                fields[field_name] = value
            continue

        field_name = ALIAS_LOOKUP.get(_normalize_label(line))
        if not field_name or field_name in fields:
            continue

        for next_line in lines[index + 1 :]:
            if ALIAS_LOOKUP.get(_normalize_label(next_line)):
                break
            split_next = _split_label_value(next_line)
            if split_next is not None and ALIAS_LOOKUP.get(_normalize_label(split_next[0])):
                break
            value = _clean_text(next_line)
            if value:
                fields[field_name] = value
                break

    return fields


def _parse_number(value: str) -> float:
    match = re.search(r"-?\d+(?:,\d{3})*(?:\.\d+)?|-?\d+(?:\.\d+)?", value or "")
    if not match:
        return 0.0
    try:
        return float(match.group(0).replace(",", ""))
    except ValueError:
        return 0.0


def _format_number(value: float) -> str:
    return str(int(value)) if float(value).is_integer() else f"{value:.2f}".rstrip("0").rstrip(".")


def _build_item_description(fields: dict[str, str]) -> str:
    comments = fields.get("comments", "")
    if comments:
        return comments
    return "Shipment Request Form freight"


def _build_comments(fields: dict[str, str]) -> str:
    return fields.get("comments", "")


def _validate_doc_upload_fields(fields: dict[str, str]) -> None:
    missing: list[str] = []
    required_fields = (
        ("Delivery Facility", "delivery_facility"),
        ("Delivery Street Address", "delivery_street_address"),
        ("Delivery City/State/Zip", "delivery_city_state_zip"),
        ("Pallet Qty", "pallet_qty"),
    )
    for label, field_name in required_fields:
        if not fields.get(field_name, "").strip():
            missing.append(label)

    if not any(
        fields.get(field_name, "").strip()
        for field_name in ("delivery_number", "kkg_po", "carrier_pro")
    ):
        missing.append("Delivery #, KKG PO#, or Carrier Pro")

    if missing:
        raise ValueError(
            "DOCX is missing required minimum fields for BOL generation: "
            + ", ".join(missing)
            + "."
        )


def _build_standard_record(fields: dict[str, str]) -> BolStandardRecord:
    origin_facility = fields.get("origin_facility", "")
    origin_street = fields.get("origin_street_address", "")
    origin_city_state_zip = fields.get("origin_city_state_zip", "")
    delivery_number = fields.get("delivery_number", "")
    kkg_po = fields.get("kkg_po", "")
    carrier_pro = fields.get("carrier_pro", "")
    bol_number = delivery_number or kkg_po or carrier_pro
    kkg_load = fields.get("kkg_load", "") or "1"
    pallet_qty = fields.get("pallet_qty", "")
    raw_pallet_weight = fields.get("pallet_weight", "")
    pallet_weight = _format_number(_parse_number(raw_pallet_weight)) if raw_pallet_weight else ""
    item_description = _build_item_description(fields)

    ship_from = BolAddressBlock(
        company=origin_facility,
        street=origin_street,
        city_state_zip=origin_city_state_zip,
        attn="",
    )
    bill_to = BolAddressBlock(
        company="Trident Transport, LLC",
        street="505 Riverfront Pkwy",
        city_state_zip="Chattanooga, TN 37402",
        attn="",
    )

    return BolStandardRecord(
        bol_number=bol_number,
        ship_date="",
        carrier=fields.get("carrier_scac", ""),
        kk_load_number=kkg_load,
        kk_po_number=kkg_po,
        po_number=kkg_po,
        dc_number=delivery_number,
        consignee_company=fields.get("delivery_facility", ""),
        consignee_street=fields.get("delivery_street_address", ""),
        consignee_city_state_zip=fields.get("delivery_city_state_zip", ""),
        ship_from=ship_from,
        bill_to=bill_to,
        seal_number_blank="",
        comments=_build_comments(fields),
        item_lines=[
            BolStandardItemLine(
                source_row_number=1,
                pallet_qty=pallet_qty,
                type="PLT",
                po_number=kkg_po,
                item_description=item_description,
                item_number=fields.get("item_number", ""),
                upc=fields.get("upc", ""),
                skids=pallet_qty,
                weight_each=pallet_weight,
            )
        ],
        total_skids=_parse_number(pallet_qty),
        is_ready=True,
        status="Ready",
        selected_for_generation=True,
        missing_required_fields=[],
        warnings=[],
        generation_skip_reason=None,
        conversion_skip_reason=None,
        issues=[],
        carrier_pro_number=carrier_pro,
    )


def parse_bol_doc_upload(file: Any) -> BolDocUploadParseResult:
    """Parse a Shipment Request Form DOCX into one Standard-family BOL record."""

    if file is None:
        raise ValueError("No DOCX file uploaded. Upload a Shipment Request Form DOCX to parse.")

    filename = str(getattr(file, "name", "") or "")
    if filename and not filename.lower().endswith(".docx"):
        raise ValueError("Uploaded file is not a DOCX file. Upload a .docx Shipment Request Form.")

    try:
        file.seek(0)
        doc = Document(file)
    except Exception as exc:
        raise ValueError(f"Unable to read DOCX upload: {exc}") from exc

    lines = _extract_document_lines(doc)
    readable_text = "\n".join(lines).strip()
    if not readable_text:
        raise ValueError("DOCX contains no readable text.")

    fields = _extract_fields_from_lines(lines)
    _validate_doc_upload_fields(fields)
    record = _build_standard_record(fields)

    return BolDocUploadParseResult(
        fields=fields,
        records=[record],
        readable_text=readable_text,
    )
