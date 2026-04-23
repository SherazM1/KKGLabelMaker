"""DOCX generation service for Multistop-mode BOL records."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from tempfile import mkdtemp
from xml.sax.saxutils import escape
import zipfile

from docx import Document
from docx.table import Table

from app.models.bol_multistop_record import BolMultistopRecord
from app.services.bol_standard_docx_generator import (
    DocxGenerationNotice,
    FailedDocxRecord,
    GeneratedDocxFile,
    SkippedDocxRecord,
    StandardDocxGenerationResult,
)
from app.utils.bol_facilities import BolFacilityRecord


MULTISTOP_TEMPLATE_PATH = Path("app/templates/multistop_bol_template.docx")
LEFT_MERGE = "\u00ab"
RIGHT_MERGE = "\u00bb"


def _tok(name: str) -> str:
    return f"{LEFT_MERGE}{name}{RIGHT_MERGE}"


def _sanitize_filename_part(value: str) -> str:
    cleaned = "".join(char if char.isalnum() or char in ("-", "_") else "_" for char in value)
    cleaned = cleaned.strip("_")
    return cleaned or "unknown"


def _unique_destination_path(directory: Path, base_name: str, extension: str) -> Path:
    candidate = directory / f"{base_name}{extension}"
    if not candidate.exists():
        return candidate

    suffix = 2
    while True:
        candidate = directory / f"{base_name}_{suffix}{extension}"
        if not candidate.exists():
            return candidate
        suffix += 1


def _format_number(value: float) -> str:
    return str(int(value)) if float(value).is_integer() else f"{value:.2f}".rstrip("0").rstrip(".")


def _format_ship_date_for_template(raw_ship_date: str) -> str:
    value = (raw_ship_date or "").strip()
    if not value:
        return ""

    normalized = value.replace("T", " ")
    for sep in (" ", "."):
        if sep in normalized:
            date_candidate = normalized.split(sep, 1)[0].strip()
            if date_candidate and any(char.isdigit() for char in date_candidate):
                normalized = date_candidate
                break

    parse_formats = (
        "%Y-%m-%d",
        "%m/%d/%Y",
        "%m/%d/%y",
        "%m-%d-%Y",
        "%m-%d-%y",
    )
    for parse_format in parse_formats:
        try:
            parsed = datetime.strptime(normalized, parse_format)
            return parsed.strftime("%m/%d/%Y")
        except ValueError:
            continue

    try:
        parsed_iso = datetime.fromisoformat(value.replace("Z", ""))
        return parsed_iso.strftime("%m/%d/%Y")
    except ValueError:
        return normalized


def _replace_text_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    text_nodes = paragraph._p.findall(".//w:t", paragraph._p.nsmap)
    instr_nodes = paragraph._p.findall(".//w:instrText", paragraph._p.nsmap)
    for node in [*text_nodes, *instr_nodes]:
        text = node.text or ""
        updated = text
        for source, target in replacements.items():
            if source in updated:
                updated = updated.replace(source, target)
        if updated != text:
            node.text = updated


def _replace_text_in_document(
    doc: Document, replacements: dict[str, str], *, include_xml_tree: bool = True
) -> None:
    def _replace_in_table_collection(tables: list[Table]) -> None:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_text_in_paragraph(paragraph, replacements)
                    _replace_in_table_collection(cell.tables)

    def _replace_in_element_tree(element) -> None:
        text_nodes = element.findall(".//w:t", element.nsmap)
        instr_nodes = element.findall(".//w:instrText", element.nsmap)
        for node in [*text_nodes, *instr_nodes]:
            text = node.text or ""
            updated = text
            for source, target in replacements.items():
                if source in updated:
                    updated = updated.replace(source, target)
            if updated != text:
                node.text = updated

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)

    _replace_in_table_collection(doc.tables)
    if include_xml_tree:
        _replace_in_element_tree(doc.element)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)
        _replace_in_table_collection(section.header.tables)
        if include_xml_tree:
            _replace_in_element_tree(section.header._element)

        for paragraph in section.footer.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)
        _replace_in_table_collection(section.footer.tables)
        if include_xml_tree:
            _replace_in_element_tree(section.footer._element)


def _resolve_comment_for_record(record_comment: str, batch_comment: str | None) -> str:
    record_value = (record_comment or "").strip()
    if record_value:
        return record_value
    return (batch_comment or "").strip()


def _postprocess_comments_in_saved_docx(destination: Path, resolved_comment: str) -> bool:
    xml_path = "word/document.xml"
    with zipfile.ZipFile(destination, "r") as archive:
        if xml_path not in archive.namelist():
            return False
        file_payloads = {name: archive.read(name) for name in archive.namelist()}

    xml_text = file_payloads[xml_path].decode("utf-8", errors="ignore")
    updated_xml = xml_text
    safe_comment = escape(resolved_comment)

    if resolved_comment:
        updated_xml = updated_xml.replace("Comments:</w:t>", f"Comments: {safe_comment}</w:t>", 1)
        updated_xml = updated_xml.replace("COMMENTS:</w:t>", f"COMMENTS: {safe_comment}</w:t>", 1)

    comment_label_populated = bool(
        resolved_comment
        and (
            f"Comments: {safe_comment}</w:t>" in updated_xml
            or f"COMMENTS: {safe_comment}</w:t>" in updated_xml
        )
    )

    if updated_xml == xml_text:
        return comment_label_populated

    file_payloads[xml_path] = updated_xml.encode("utf-8")
    with zipfile.ZipFile(destination, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for name, payload in file_payloads.items():
            archive.writestr(name, payload)

    return comment_label_populated


def _populate_ship_from_block(doc: Document, selected_facility: BolFacilityRecord) -> bool:
    name_value = selected_facility["facility_name"]
    street_value = selected_facility["address"]
    city_state_zip_value = selected_facility["location"]

    for table in doc.tables:
        in_ship_from_block = False
        for row in table.rows:
            row_cells = row.cells
            if not row_cells:
                continue

            row_text_upper = " ".join(cell.text.strip() for cell in row_cells).upper()
            first_cell_text = row_cells[0].text.strip().upper().replace(" ", "")

            if "FROM (SHIPPER)" in row_text_upper:
                in_ship_from_block = True
                continue
            if in_ship_from_block and "TO (CONSIGNEE)" in row_text_upper:
                return True
            if not in_ship_from_block:
                continue

            if first_cell_text == "NAME" and len(row_cells) > 1:
                row_cells[1].text = name_value
            elif first_cell_text == "STREET" and len(row_cells) > 1:
                row_cells[1].text = street_value
            elif first_cell_text in {"CITY/ST/ZIP", "CITY/STATE/ZIP"} and len(row_cells) > 1:
                row_cells[1].text = city_state_zip_value

    return False


def _template_replacements(record: BolMultistopRecord) -> dict[str, str]:
    replacements = {
        _tok("BOL_"): record.bol_number,
        _tok("ship_date"): _format_ship_date_for_template(record.ship_date),
        _tok("Carrier"): record.carrier,
        _tok("load"): record.load_number,
        _tok("KK_PO"): record.kk_po_number,
        _tok("KK_Load"): record.kk_load_number,
        _tok("DELIVERY_1_DC"): record.delivery_1_dc,
        _tok("DELIVERY_1_ADDRESS"): record.delivery_1_address,
        _tok("DELIVERY_2_DC"): record.delivery_2_dc,
        _tok("DELIVERY_2_ADDRESS"): record.delivery_2_address,
        _tok("DELIVERY_3_DC"): record.delivery_3_dc,
        _tok("DELIVERY_3_ADDRESS"): record.delivery_3_address,
        _tok("DC_1"): record.dc_1,
        _tok("CASE_1"): record.case_1,
        _tok("PO_1"): record.po_1,
        _tok("Pallet_Description_1"): record.pallet_description_1,
        _tok("PLT_1"): record.plt_1,
        _tok("WEIGHT_1"): record.weight_1,
        _tok("DC_2"): record.dc_2,
        _tok("CASE_2"): record.case_2,
        _tok("PO_2"): record.po_2,
        _tok("Pallet_Description_2"): record.pallet_description_2,
        _tok("PLT_2"): record.plt_2,
        _tok("WEIGHT_2"): record.weight_2,
        _tok("DC_3"): record.dc_3,
        _tok("CASE_3"): record.case_3,
        _tok("PO_3"): record.po_3,
        _tok("Pallet_Description_3"): record.pallet_description_3,
        _tok("PLT_3"): record.plt_3,
        _tok("WEIGHT_3"): record.weight_3,
        _tok("Total_Case"): _format_number(record.total_case),
        _tok("Total_Pallet"): _format_number(record.total_pallet),
        _tok("Total_Ship_Weight"): _format_number(record.total_ship_weight),
    }
    return replacements


def generate_multistop_docx_set(
    records: list[BolMultistopRecord],
    selected_facility: BolFacilityRecord | None,
    batch_comment: str | None = None,
    template_path: Path | None = None,
    output_dir: Path | None = None,
    file_name_prefix: str = "multistop_bol",
) -> StandardDocxGenerationResult:
    if selected_facility is None:
        raise ValueError(
            "No ship-from facility is selected. Select a facility in BOL Generator before DOCX generation."
        )

    resolved_template = template_path or MULTISTOP_TEMPLATE_PATH
    if not resolved_template.exists():
        raise FileNotFoundError(f"Template file not found: {resolved_template}")

    output_root = output_dir or Path(mkdtemp(prefix="kkg_multistop_bol_docx_"))
    output_root.mkdir(parents=True, exist_ok=True)

    generated: list[GeneratedDocxFile] = []
    skipped: list[SkippedDocxRecord] = []
    failed: list[FailedDocxRecord] = []
    notices: list[DocxGenerationNotice] = []

    for record in records:
        bol_label = record.bol_number or "(missing BOL #)"
        record.generation_skip_reason = None

        if not record.selected_for_generation:
            reason = "Record excluded in review."
            record.generation_skip_reason = reason
            skipped.append(SkippedDocxRecord(bol_number=bol_label, reason=reason))
            continue

        if record.stop_count > 3:
            reason = "Unsupported stop count: more than 3 stops."
            record.generation_skip_reason = reason
            skipped.append(SkippedDocxRecord(bol_number=bol_label, reason=reason))
            continue

        if not record.is_ready:
            reason = "Record is not ready for DOCX generation."
            if record.status == "Unsupported Stop Count":
                reason = "Unsupported stop count: more than 3 stops."
            elif record.missing_required_fields:
                reason = "Missing required data: " + ", ".join(record.missing_required_fields)
            elif record.issues:
                reason = "; ".join(record.issues)
            record.generation_skip_reason = reason
            skipped.append(SkippedDocxRecord(bol_number=bol_label, reason=reason))
            continue

        try:
            doc = Document(str(resolved_template))
            replacements = _template_replacements(record)
            _replace_text_in_document(doc, replacements, include_xml_tree=True)
            ship_from_populated = _populate_ship_from_block(doc, selected_facility)
            if not ship_from_populated:
                notices.append(
                    DocxGenerationNotice(
                        bol_number=bol_label,
                        message="Could not confirm ship-from block location in template.",
                    )
                )

            safe_bol = _sanitize_filename_part(record.bol_number)
            safe_load = _sanitize_filename_part(record.load_number)
            destination = _unique_destination_path(
                output_root, f"{file_name_prefix}_{safe_bol}_{safe_load}", ".docx"
            )
            filename = destination.name
            doc.save(str(destination))

            resolved_comment = _resolve_comment_for_record(record.comments, batch_comment)
            comment_label_populated = _postprocess_comments_in_saved_docx(
                destination, resolved_comment
            )
            if resolved_comment and not comment_label_populated:
                notices.append(
                    DocxGenerationNotice(
                        bol_number=bol_label,
                        message=(
                            "Resolved comment was non-empty but could not be confirmed "
                            "at the visible Comments label in word/document.xml."
                        ),
                    )
                )

            generated.append(
                GeneratedDocxFile(
                    bol_number=bol_label,
                    file_name=filename,
                    file_path=str(destination.resolve()),
                )
            )
        except Exception as exc:
            failed.append(FailedDocxRecord(bol_number=bol_label, error=str(exc)))

    if not generated and not failed:
        raise ValueError("No selected and ready records are available for DOCX generation.")

    return StandardDocxGenerationResult(
        output_dir=str(output_root.resolve()),
        generated_files=generated,
        skipped_records=skipped,
        failed_records=failed,
        notices=notices,
    )

