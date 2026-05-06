"""DOCX generator for SKID tags."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.models.skid_tag import SkidTag


TIMES_NEW_ROMAN = "Times New Roman"
APTOS_NARROW = "Aptos Narrow"
TAG_FONT_SIZE = Pt(72)
PO_VALUE_FONT_SIZE = Pt(60)
LINE_HEIGHT = Pt(76)
LINE_SPACING = Pt(12)


def _set_run_font(run, font_name: str, font_size: Pt) -> None:
    run.font.name = font_name
    run.font.size = font_size
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)


def _add_centered_paragraph(
    document: Document,
    text: str,
    space_after: Pt = LINE_SPACING,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = space_after
    paragraph.paragraph_format.line_spacing = LINE_HEIGHT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = paragraph.add_run(text)
    _set_run_font(run, TIMES_NEW_ROMAN, TAG_FONT_SIZE)
    return paragraph


def _add_po_paragraph(document: Document, po_value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = LINE_SPACING
    paragraph.paragraph_format.line_spacing = LINE_HEIGHT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    label_run = paragraph.add_run("PO ")
    _set_run_font(label_run, TIMES_NEW_ROMAN, TAG_FONT_SIZE)

    value_run = paragraph.add_run(po_value)
    _set_run_font(value_run, APTOS_NARROW, PO_VALUE_FONT_SIZE)


def _add_tag_page(document: Document, tag: SkidTag):
    _add_centered_paragraph(document, f"DC {tag.dc}")
    _add_po_paragraph(document, tag.po_display)
    _add_centered_paragraph(document, f"UPC {tag.upc}")
    _add_centered_paragraph(document, tag.pallet_display)
    return _add_centered_paragraph(document, f"Qty: {tag.quantity}", space_after=Pt(0))


def generate_skid_tags_docx(tags: list[SkidTag]) -> bytes:
    if not tags:
        raise ValueError("No SKID tags provided for DOCX generation.")

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    for index, tag in enumerate(tags):
        last_paragraph = _add_tag_page(document, tag)
        if index < len(tags) - 1:
            last_paragraph.add_run().add_break(WD_BREAK.PAGE)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()
